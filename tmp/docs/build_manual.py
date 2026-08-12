from __future__ import annotations

import os
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "doc"
ASSETS = ROOT / "docs" / "assets" / "bs-manual"
OUT.mkdir(parents=True, exist_ok=True)


def git_commit() -> str:
    try:
        return subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "未知"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:%s" % edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in kwargs[edge]:
                    element.set(qn("w:%s" % key), str(kwargs[edge][key]))


def set_cell_text(cell, text: str, *, bold=False, color="1F2937", size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        set_cell_text(hdr.cells[i], value, bold=True, color="FFFFFF", size=9)
        set_cell_shading(hdr.cells[i], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.8)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], "F2F6FA")
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(paragraph):
    """Insert an updateable Word table-of-contents field with a visible fallback result."""
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    begin_run._r.extend([fld_begin, instr, fld_sep])
    placeholder = paragraph.add_run("打开 Microsoft Word 后更新目录")
    placeholder.font.color.rgb = RGBColor(127, 127, 127)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(fld_end)


def enable_update_fields(doc):
    settings = doc.settings._element
    for old in settings.findall(qn("w:updateFields")):
        settings.remove(old)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_code(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    set_cell_border(cell, top={"val": "single", "sz": 4, "color": "CBD5E1"}, bottom={"val": "single", "sz": 4, "color": "CBD5E1"}, left={"val": "single", "sz": 4, "color": "CBD5E1"}, right={"val": "single", "sz": 4, "color": "CBD5E1"})
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for n, line in enumerate(text.splitlines()):
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
        if n < len(text.splitlines()) - 1:
            r.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_note(doc, title: str, text: str, fill="EAF3F8"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, top={"val": "single", "sz": 5, "color": "5B9BD5"}, bottom={"val": "single", "sz": 5, "color": "5B9BD5"}, left={"val": "single", "sz": 12, "color": "2F75B5"}, right={"val": "single", "sz": 5, "color": "5B9BD5"})
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{title}：")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.size = Pt(9.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_step(doc, number, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}. {title}：")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(text)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(89, 89, 89)


def add_image(doc, filename, caption, width_cm=15.5):
    path = ASSETS / filename
    if not path.exists():
        add_note(doc, "截图缺失", f"未找到 {filename}，请在 docs/assets/bs-manual 中补齐。", "FCE4D6")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_module(doc, *, title, purpose, preconditions, params, steps, results, errors, screenshot=None, result_images=None):
    doc.add_heading(title, level=2)
    doc.add_paragraph(purpose)
    doc.add_heading("使用前准备", level=3)
    add_bullets(doc, preconditions)
    doc.add_heading("参数说明", level=3)
    add_table(doc, ["参数/区域", "说明"], params, widths=[4.2, 12.0])
    doc.add_heading("操作步骤", level=3)
    for n, (st, tx) in enumerate(steps, 1):
        add_step(doc, n, st, tx)
    doc.add_heading("结果解读", level=3)
    add_bullets(doc, results)
    if result_images:
        for fn, cap in result_images:
            add_image(doc, fn, cap, 13.5)
    if screenshot:
        add_image(doc, screenshot[0], screenshot[1], 15.5)
    doc.add_heading("常见问题", level=3)
    add_bullets(doc, errors)


def add_new_page(doc):
    """Insert a section break that survives Word, LibreOffice and Pandoc PDF conversion."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header.is_linked_to_previous = True
    sec.footer.is_linked_to_previous = True
    return sec


def build():
    doc = Document()
    enable_update_fields(doc)
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [("Title", 25, "1F4E79"), ("Heading 1", 17, "1F4E79"), ("Heading 2", 13, "2F75B5"), ("Heading 3", 11, "404040")]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    if "Caption" in doc.styles:
        doc.styles["Caption"].font.name = "Microsoft YaHei"

    # Header/footer.
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("流程行业动态风险管控工具集 V1.0  ·  软件使用手册")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(127, 127, 127)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("内部使用  |  第 ")
    fr.font.size = Pt(8)
    add_field(footer, "PAGE")
    fr2 = footer.add_run(" 页")
    fr2.font.size = Pt(8)

    # Cover.
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("流程行业动态风险管控工具集")
    r.font.size = Pt(27)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("V1.0 软件使用手册")
    r.font.size = Pt(23)
    r.bold = True
    r.font.color.rgb = RGBColor(47, 117, 181)
    doc.add_paragraph()
    cover = add_table(doc, ["项目", "内容"], [
        ("软件形态", "BS 本机 Web 版（浏览器访问）"),
        ("适用对象", "普通用户"),
        ("适用版本", "V1.0；构建提交 " + git_commit()),
        ("编制日期", str(date.today())),
    ], widths=[4.0, 12.2])
    doc.add_paragraph()
    add_note(doc, "使用范围", "本手册用于指导普通用户访问本机 Web 版软件、完成各功能模块操作并查看分析结果。", "E2F0D9")
    add_new_page(doc)

    # TOC and document control.
    doc.add_heading("目录", level=1)
    toc = doc.add_paragraph()
    add_toc(toc)
    add_note(doc, "目录更新", "右键目录，选择“更新域”，再选择“更新整个目录”。", "FFF2CC")
    add_new_page(doc)

    doc.add_heading("1 文档说明", level=1)
    doc.add_heading("1.1 文档目的", level=2)
    doc.add_paragraph("本手册用于指导普通用户完成软件启动与访问、功能模块切换、参数填写、任务执行、结果查看以及常见问题处理。")
    doc.add_heading("1.2 修订记录", level=2)
    add_table(doc, ["版本", "日期", "说明"], [("V1.0", str(date.today()), "基于当前 BS 本机 Web 版本编制；记录提交 " + git_commit()), ("-", "-", "后续修订请补充变更范围和验证结果。")], widths=[3, 3, 10.2])
    doc.add_heading("1.3 术语说明", level=2)
    add_table(doc, ["术语", "含义"], [
        ("本机 Web 版", "在当前计算机启动后，通过浏览器访问的软件版本。"),
        ("任务", "用户提交的一次训练、仿真或分析操作。"),
        ("异常检测概率图", "用于展示异常检测概率变化的结果图。"),
        ("PFDavg", "平均要求时失效概率，用于在线 SIL 验证结果判读。"),
    ], widths=[4, 12.2])

    doc.add_heading("2 软件简介与使用准备", level=1)
    doc.add_paragraph("流程行业动态风险管控工具集通过浏览器提供异常检测、风险分析、控制模型训练、控制仿真、SDG-HAZOP 和在线 SIL 验证等功能。普通用户可在统一界面中设置参数、提交任务并查看图表结果。")
    add_table(doc, ["功能类别", "主要用途"], [
        ("异常行为检测", "查看网络拓扑和异常检测概率。"),
        ("风险动态分析", "完成威胁分类、风险评分和风险场景动态匹配。"),
        ("风险管控优化决策", "完成控制模型训练评估和优化控制仿真验证。"),
        ("SIS 自主化检测", "配置 SDG 模型并查看 HAZOP 分析及推荐结果。"),
        ("在线 SIL 验证", "计算 PFDavg、SIL 等级并查看概率分布。"),
    ], widths=[5, 11.2])
    doc.add_heading("2.1 使用前准备", level=2)
    add_bullets(doc, [
        "确认软件运行环境已由技术支持人员配置完成。",
        "建议使用 Microsoft Edge 或 Google Chrome 浏览器。",
        "准备需要导入的数据文件或模型文件；首次使用时可先采用页面默认参数。",
        "运行耗时任务前关闭不必要的软件，并保持浏览器窗口开启。",
    ])
    doc.add_heading("2.2 简易启动与访问", level=2)
    add_step(doc, 1, "切换到项目目录", "如果 PowerShell 当前显示的是 C:\\Users\\... 等其他目录，请先切换到项目根目录。当前部署示例为 D:\\ADRL\\Platform_sjtu。")
    add_code(doc, "Set-Location 'D:\\ADRL\\Platform_sjtu'\n& '.\\scripts\\run_web_local.ps1'")
    doc.add_paragraph("如果已经位于项目根目录，只需运行第二行；也可以从任意目录直接使用完整路径调用脚本：")
    add_code(doc, "& 'D:\\ADRL\\Platform_sjtu\\scripts\\run_web_local.ps1'")
    add_step(doc, 2, "等待启动完成", "看到“Local Web edition is ready”提示后再打开浏览器。")
    add_step(doc, 3, "访问软件", "在浏览器地址栏输入 http://127.0.0.1:8000。")
    add_step(doc, 4, "结束使用", "关闭浏览器页面；如需停止软件，在启动脚本窗口按 Ctrl+C。")
    add_note(doc, "本机访问", "当前版本仅供启动软件的这台计算机访问。若启动失败，请记录屏幕提示并联系技术支持人员。", "E2F0D9")
    add_image(doc, "01-home.png", "图 2-1  软件首页与模块导航（实际本机页面截图）", 15.5)

    doc.add_heading("3 通用界面操作", level=1)
    doc.add_heading("3.1 导航与模块切换", level=2)
    add_table(doc, ["区域", "操作要点"], [
        ("顶部导航", "点击一级模块展开下拉菜单，再点击具体功能；当前页面以高亮状态显示。"),
        ("参数面板", "输入框通常带默认值和范围限制；失焦或提交时执行边界校验。"),
        ("任务状态", "等待、运行、完成、失败由状态条和进度信息表示；训练/仿真期间不要重复提交。"),
        ("结果区", "任务完成后显示结果图、评分、曲线或文字说明。"),
        ("提示信息", "出现提示弹窗时，先按提示修正参数或文件选择，再重新提交。"),
    ], widths=[4, 12.2])
    doc.add_heading("3.2 参数填写与任务状态", level=2)
    add_bullets(doc, [
        "首次使用建议保留默认参数，确认流程正常后再逐项调整。",
        "带范围限制的数值必须在页面允许区间内，文件路径应指向可用文件。",
        "任务处于等待或运行状态时请勿重复点击执行按钮。",
        "任务失败时记录模块名称、输入参数和完整错误提示，便于技术支持人员处理。",
    ])
    doc.add_heading("3.3 刷新与重新进入", level=2)
    add_bullets(doc, [
        "仅查看已显示结果时可以刷新页面。",
        "任务仍在运行时不要关闭浏览器；如误刷新，可重新进入对应模块查看状态。",
        "页面显示异常或按钮无响应时，可先按 Ctrl+F5 强制刷新，再重新操作。",
    ])
    add_image(doc, "11-validation.png", "图 3-1  参数或运行条件提示弹窗示例", 15.5)

    add_module(doc,
        title="4.1 基于移动目标防御的异常检测",
        purpose="根据网络拓扑、攻击强度、测量噪声和过程扰动进行分析，输出网络拓扑图和异常检测概率图。",
        preconditions=["系统运行环境已由技术支持人员配置完成。", "异常检测页面可以正常打开；运行环境相关输入保持默认值。"],
        params=[("运行环境相关输入", "保持页面默认值，不建议普通用户修改。"), ("攻击强度", "默认 5-10%；允许范围 5-50%。"), ("测量噪声", "默认 2%；允许范围 1-30%。"), ("过程扰动", "默认 5%；允许范围 1-30%。"), ("结果标签", "网络拓扑图、异常检测概率图。")],
        steps=[("确认默认设置", "保持页面中的运行环境相关输入不变。"), ("设置参数", "输入攻击强度、噪声和扰动范围，保持在边界内。"), ("提交任务", "点击执行，等待状态变为完成。"), ("查看结果", "先查看拓扑图，再切换到异常检测概率图。" )],
        results=["拓扑图用于确认节点和链路结构。", "异常检测概率图用于观察不同条件下的检测概率变化。"],
        errors=["参数无法提交：检查数值是否在允许范围内。", "出现运行环境提示：保留完整弹窗内容并联系技术支持人员。", "结果图未显示：确认任务已完成后刷新页面；仍无结果时记录任务信息。"],
        screenshot=("03-anomaly-result.png", "图 4-1  异常检测参数与异常检测概率图（实际页面截图）"),
        result_images=[("result-topology.png", "图 4-2  网络拓扑图样例"), ("result-detection-probability.png", "图 4-3  异常检测概率图样例")])

    add_module(doc,
        title="4.2 潜在安全威胁识别与自动分类",
        purpose="使用分类数据集训练模型并展示混淆矩阵、识别召回率和最佳准确率。",
        preconditions=["页面中的 original、easy、hard 数据集可以正常选择。", "如使用自定义数据，应先由技术支持人员确认格式。"],
        params=[("数据集", "original、easy 或 hard。"), ("Epochs", "训练轮数，范围 1-500，默认值以页面显示为准。"), ("Batch Size", "批大小，范围 8-256。"), ("Learning Rate", "学习率，范围 0.0001-0.1。")],
        steps=[("选择数据集", "按难度或原始数据选择 original/easy/hard。"), ("填写超参数", "输入轮数、批大小和学习率。"), ("开始训练", "提交后观察任务进度，不要关闭页面。"), ("查看指标", "完成后查看混淆矩阵、召回率和准确率。" )],
        results=["混淆矩阵用于观察各类误判方向；召回率反映真实威胁被识别的比例。", "最佳准确率用于比较不同数据集或超参数组合，不应脱离测试集规模单独解读。"],
        errors=["数据集读取失败：重新选择数据集；仍失败时记录错误提示。", "参数超出范围：按页面提示修正后重新提交。", "训练失败：不要连续重复提交，记录当前参数并联系技术支持人员。"],
        screenshot=("05-classification.png", "图 4-4  分类模块输入界面（实际页面截图）"))

    add_module(doc,
        title="4.3 多评估准则融合的风险学习分析",
        purpose="对多项风险指标按权重融合，支持手工输入或随机生成指标，并以雷达图、综合评分和危险分数展示结果。",
        preconditions=["风险指标和权重的业务含义已经确认。"],
        params=[("指标值", "按页面列出的指标输入数值。"), ("权重", "为各指标配置权重；权重应与指标行对应。"), ("随机数据", "快速生成一组可复现/可比较的示例指标。"), ("结果", "雷达图、综合评分、危险分数。")],
        steps=[("配置指标", "输入指标值并检查权重。"), ("生成或评估", "可先使用随机数据，再点击评估。"), ("比较结果", "观察雷达图形状以及综合分与危险分数。" )],
        results=["雷达图适合查看指标间相对差异；综合评分用于整体比较，危险分数用于风险等级判读。", "调整权重会改变综合结果，应记录业务依据。"],
        errors=["权重和不符合要求：检查空值、负值或合计规则。", "结果为空：先生成/输入指标，再执行评估。"],
        screenshot=("04-score.png", "图 4-5  风险学习分析界面（实际页面截图）"))

    add_module(doc,
        title="4.4 风险场景动态匹配与适配方案生成算法",
        purpose="根据样例数据、时间步长、预测域和起始样本进行风险趋势分析，输出 CV、U_now、U_after 向量及适配方案。",
        preconditions=["系统已准备可用的分析数据。", "时间步长、预测域和起始样本的业务范围已经确认。"],
        params=[("时间步长", "风险趋势计算的时间间隔。"), ("预测域", "向前预测的步数/范围。"), ("起始样本", "从数据集哪个样本开始分析。"), ("CV/U 向量", "页面展示当前场景、当前控制和适配后控制向量。")],
        steps=[("选择数据", "确认页面已显示可用的样例或指定数据。"), ("配置窗口", "填写时间步长、预测域和起始样本。"), ("执行分析", "提交任务并等待结果。"), ("查看方案", "阅读风险趋势和适配方案文本。" )],
        results=["风险趋势图反映预测窗口内变化；CV、U_now、U_after 用于解释当前状态和适配动作。", "适配方案应结合工艺上下文复核，不能替代现场安全决策。"],
        errors=["数据无法读取：重新选择数据并记录提示信息。", "样本索引越界：减小起始样本或预测域。", "结果为空：确认任务已完成；仍为空时联系技术支持人员。"],
        screenshot=("06-cdq.png", "图 4-6  CDQ 参数与结果区域（实际页面截图）"))

    add_module(doc,
        title="4.5 控制模型训练评估",
        purpose="生成或读取过程控制数据，训练 DNN 模型并查看训练性能、预测误差和模型文件。",
        preconditions=["系统运行环境已由技术支持人员配置完成。", "如使用外部数据，准备符合要求的 .mat 文件。"],
        params=[("样本数", "默认 1000；允许范围 100-100000。"), ("训练轮数", "默认 50；允许范围 1-5000。"), ("隐藏层规模", "默认 64,64；最多 10 层，每层 1-4096。"), ("外部 .mat", "可选；使用前确认文件来源和格式。"), ("模型文件", "训练完成后由系统保存，供优化控制仿真使用。")],
        steps=[("确认数据", "选择系统生成数据或有效的外部数据文件。"), ("配置训练", "设置样本数、轮数和隐藏层规模。"), ("提交训练", "观察任务进度，训练过程中保持页面开启。"), ("查看结果", "确认训练性能图、预测误差图和模型生成状态。" )],
        results=["训练性能图用于观察训练过程是否趋于稳定。", "预测误差图用于查看模型预测偏差；生成的模型可供优化控制仿真使用。"],
        errors=["样本数或轮数超限：按页面范围调整。", "外部数据无法读取：重新选择文件并确认格式。", "训练中断或结果缺失：记录当前参数和错误提示并联系技术支持人员。"],
        screenshot=("07-training.png", "图 4-7  DNNTrain 参数与训练性能图（实际页面截图）"),
        result_images=[("result-training-performance.png", "图 4-8  训练性能图样例"), ("result-prediction-error.png", "图 4-9  预测误差图样例")])

    add_module(doc,
        title="4.6 优化控制仿真验证",
        purpose="加载 DNN 模型执行 MPC 仿真，输出过程轨迹、控制输入、跟踪误差和代价曲线。",
        preconditions=["先完成控制模型训练，或准备有效的 DNN 模型文件。", "确认模型文件与当前仿真任务匹配。"],
        params=[("DNN 模型", "选择训练完成后生成的模型文件。"), ("仿真时长", "范围 0.2-20 秒。"), ("预测时域", "范围 1-60。"), ("输出", "过程轨迹、控制输入、跟踪误差和代价曲线。")],
        steps=[("选择模型", "填写或选择已训练模型文件。"), ("配置仿真", "设置仿真时长和预测时域。"), ("提交 MPC", "等待任务完成。"), ("解释曲线", "联合观察轨迹、控制输入、跟踪误差和代价。" )],
        results=["轨迹图用于判断状态跟踪；控制输入用于检查执行器动作；跟踪误差和代价曲线用于比较控制品质。", "仿真结果应结合模型训练数据和工艺约束一起复核。"],
        errors=["模型无法读取：重新选择正确的模型文件。", "仿真参数越界：按页面范围修正。", "结果图缺失：确认任务已完成；仍缺失时记录提示并联系技术支持人员。"],
        screenshot=("08-mpc.png", "图 4-10  MPC 仿真界面（实际页面截图）"),
        result_images=[("result-process-control-trajectory.png", "图 4-11  过程控制轨迹样例"), ("result-control-input.png", "图 4-12  控制输入样例"), ("result-tracking-error.png", "图 4-13  跟踪误差样例"), ("result-cost-curve.png", "图 4-14  代价曲线样例")])

    add_module(doc,
        title="4.7 SDG-HAZOP",
        purpose="配置 SDG 节点和边，结合概率与模糊术语开展 HAZOP 分析，并生成 SIS 推荐结果。",
        preconditions=["页面可加载示例配置；如使用自定义场景，准备节点、边及术语参数。"],
        params=[("节点/边", "描述工艺节点及其连接关系。"), ("概率", "输入事件/状态概率。"), ("模糊术语", "使用页面支持的模糊语言变量。"), ("分析结果", "风险分析、关键路径和 SIS 推荐。")],
        steps=[("加载示例", "先使用示例配置确认页面和图形正常。"), ("编辑图模型", "增加/删除节点和边，保持关系闭合。"), ("配置术语", "填写概率和模糊项。"), ("执行分析", "查看风险结果与推荐说明。" )],
        results=["图形区域用于检查结构；分析区用于查看 HAZOP 风险项和 SIS 建议。", "推荐结果需由工艺和功能安全人员复核。"],
        errors=["节点/边格式错误：检查 ID 唯一性和引用关系。", "概率或术语缺失：补齐必填配置后再分析。"],
        screenshot=("09-sdg.png", "图 4-15  SDG-HAZOP 图模型与配置（实际页面截图）"))

    add_module(doc,
        title="4.8 在线 SIL 验证",
        purpose="基于 GSPN-MC 模型，对 M/N 表决结构和共因失效参数进行蒙特卡洛仿真，输出 PFDavg、SIL 结果和概率分布图。",
        preconditions=["确认仿真参数符合项目安全分析口径。", "任务运行期间保持浏览器页面开启；大规模仿真可能耗时较长。"],
        params=[("M/N", "M/N 表决结构，页面范围 1-10。"), ("失效率", "可直接输入 lambda（FIT）或选择估算。"), ("TI/MRT", "测试间隔、平均修复时间。"), ("Total/Partial beta", "总共因失效和部分共因失效比例。"), ("仿真", "默认 nsim=500、years=10000；按页面范围调整。")],
        steps=[("设置结构", "填写 M、N 以及失效率。"), ("设置维护参数", "填写 TI、MRT 和共因失效参数。"), ("设置仿真规模", "配置仿真次数和仿真年数。"), ("运行并判读", "提交后查看 PFDavg、SIL 和分布图。" )],
        results=["PFDavg 是核心量化结果；SIL 结果用于与目标等级比较。", "概率分布图用于观察仿真离散性；应同时记录输入参数和随机仿真规模。"],
        errors=["M/N 不合法：确保 1 <= M <= N <= 10。", "共因失效比例不合理：检查 Total/Partial 参数和合计规则。", "任务超时：先降低 nsim/years 验证参数，再逐步提高规模。"],
        screenshot=("10-sil.png", "图 4-16  在线 SIL 参数界面（实际页面截图）"))

    doc.add_heading("5 结果查看与保存", level=1)
    add_table(doc, ["结果类型", "查看与保存建议"], [
        ("结果图", "使用页面中的结果标签切换图片；如需留存，可使用页面提供的下载功能或浏览器截图。"),
        ("评分与数值", "记录任务参数、综合评分、准确率、PFDavg 或 SIL 等关键数值。"),
        ("训练模型", "确认页面提示模型生成成功；后续仿真时选择与本次训练对应的模型。"),
        ("文字方案", "复制或截图保存风险趋势、适配方案和 SIS 推荐，并标注任务日期。"),
    ], widths=[4, 12.2])
    doc.add_heading("5.1 结果留存建议", level=2)
    add_bullets(doc, [
        "保存结果时同时记录模块名称、输入参数和执行日期。",
        "比较多组结果时使用一致的参数口径，并为每组结果添加清晰名称。",
        "训练模型和仿真结果应成组保存，避免误选其他训练任务生成的模型。",
        "涉及安全结论的结果应由相关专业人员复核后使用。",
    ])

    doc.add_heading("6 常见问题处理", level=1)
    add_table(doc, ["现象", "普通用户处理方法"], [
        ("无法打开软件页面", "确认一键启动窗口仍在运行，重新输入 http://127.0.0.1:8000；仍无法访问时记录启动窗口提示。"),
        ("参数无法提交", "检查是否存在空值、非数字字符或超出范围的数值，并按页面提示修正。"),
        ("数据或模型文件无法读取", "重新选择正确文件，确认文件未被移动、改名或由其他程序占用。"),
        ("任务长时间停留在等待或运行状态", "不要重复提交；等待一段时间后记录模块、参数和当前进度。"),
        ("任务显示失败", "完整记录错误弹窗、任务参数和操作步骤，然后联系技术支持人员。"),
        ("结果图未显示", "确认任务状态为完成，切换结果标签或按 Ctrl+F5 刷新页面后重试。"),
        ("页面显示旧内容或布局异常", "关闭旧页面，重新打开软件地址，并按 Ctrl+F5 强制刷新。"),
        ("误关闭或刷新页面", "重新进入对应模块；如任务状态无法恢复，记录原参数后重新提交。"),
    ], widths=[5.2, 11.0])
    add_note(doc, "联系技术支持", "提交问题时请提供软件版本、功能模块、输入参数、操作时间和完整错误截图。", "E2F0D9")

    doc.add_heading("附录 A 主要模块速查", level=1)
    add_table(doc, ["模块", "入口", "主要输出"], [
        ("异常检测", "异常行为检测 → 基于移动目标防御的异常检测", "网络拓扑图、异常检测概率图"),
        ("分类", "风险动态分析 → 潜在安全威胁识别与自动分类", "混淆矩阵、召回率、准确率"),
        ("评分", "风险动态分析 → 多评估准则融合的风险学习分析", "雷达图、综合评分、危险分数"),
        ("CDQ", "风险动态分析 → 风险场景动态匹配与适配方案生成算法", "风险趋势、CV/U 向量、适配方案"),
        ("DNNTrain", "风险管控优化决策 → 控制模型训练评估", "训练性能、预测误差、模型文件"),
        ("MPC", "风险管控优化决策 → 优化控制仿真验证", "轨迹、控制输入、跟踪误差、代价"),
        ("SDG-HAZOP", "SIS自主化检测 → SDG-HAZOP", "风险分析、SIS 推荐"),
        ("在线 SIL", "在线SIL验证 → 基于GSPN-MC模型的动态化SIL验证方法", "PFDavg、SIL、概率分布图"),
    ], widths=[3.5, 8.2, 4.5])
    doc.add_heading("附录 B 界面截图说明", level=1)
    doc.add_paragraph("本手册中的界面截图均来自当前本机 Web 版实际页面，结果图使用项目现有真实样例。界面统一使用“异常检测概率图”这一业务名称。")

    # Ensure all paragraphs use a readable default font where a style did not set it.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(10.5)
            run.font.name = run.font.name or "Microsoft YaHei"
            if run._element.rPr is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name or "Microsoft YaHei")
    doc.core_properties.title = "流程行业动态风险管控工具集 V1.0 软件使用手册"
    doc.core_properties.subject = "BS 本机 Web 版普通用户操作手册"
    doc.core_properties.author = "项目组"
    doc.core_properties.comments = f"构建提交 {git_commit()}"
    target = OUT / "流程行业动态风险管控工具集V1.0软件使用手册.docx"
    doc.save(target)
    print(target)


if __name__ == "__main__":
    build()
