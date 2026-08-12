from pathlib import Path
from zipfile import ZipFile

from docx import Document


matches = list(Path("output/doc").glob("*.docx"))
assert len(matches) == 1, matches
path = matches[0]
document = Document(path)
parts = [paragraph.text for paragraph in document.paragraphs]
for table in document.tables:
    for row in table.rows:
        parts.extend(cell.text for cell in row.cells)
text = "\n".join(parts)

forbidden = [
    "普通用户、系统管理员",
    "帮助管理员",
    "pip install",
    "MCR_ROOT",
    "mclmcrrt24_2.dll",
    "npm run",
    "/api/",
    "FastAPI",
    "TaskManager",
    "gridattackpkg",
    "dnnmpcpkg",
    "E:\\MATLAB2024",
    "D:\\ana3\\envs",
]
hits = {term: text.count(term) for term in forbidden if term in text}
assert not hits, hits
assert "适用对象\n普通用户" in text
assert text.count("技术支持人员") >= 3

headings = [
    (paragraph.style.name, paragraph.text)
    for paragraph in document.paragraphs
    if paragraph.style.name.startswith("Heading")
]
module_headings = [title for _, title in headings if title.startswith("4.")]
assert len(module_headings) == 8, module_headings
assert len(document.inline_shapes) == 18

with ZipFile(path) as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")
    settings_xml = archive.read("word/settings.xml").decode("utf-8")
assert 'TOC \\o "1-3" \\h \\z \\u' in document_xml
assert "w:updateFields" in settings_xml
assert 'w:val="true"' in settings_xml or 'w:val="1"' in settings_xml

print("forbidden_hits=", hits)
print("heading_count=", len(headings))
print("module_headings=", module_headings)
print("support_mentions=", text.count("技术支持人员"))
print("toc_field= True")
print("update_fields= True")
print("inline_shapes=", len(document.inline_shapes))
